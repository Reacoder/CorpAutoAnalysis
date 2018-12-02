#!/usr/local/bin/python
# coding: utf-8

import os

from docx import Document
from docx.shared import Inches

import config
import util

document = Document()


def generate_doc():
    code = config.CODE
    my_dir = 'private/' + code + '/'
    my_dict = util.load_json(code)
    pic_width = Inches(7)
    # 标题
    document.add_heading(my_dict['code_name'], 0)
    # 走势图
    document.add_heading('走势图', level=1)
    document.add_heading('1.周k线', level=2)
    document.add_picture(my_dir + 'img_trend_week.png', width=pic_width)
    document.add_heading('2.日k线', level=2)
    document.add_picture(my_dir + 'img_trend_day.png', width=pic_width)
    # 一.公司概要
    document.add_heading('一.公司概要', level=1)
    document.add_picture(my_dir + 'img_brief.png', width=pic_width)
    # 二.业务构成
    document.add_heading('二.业务构成', level=1)
    document.add_picture(my_dir + 'img_operate_product.png', width=pic_width)
    if os.path.isfile(my_dir + 'img_costs_gross_profit_rate.png'):
        document.add_heading('1.毛利率', level=2)
        document.add_paragraph(my_dict['gross_profit_rate'].replace('\n', ' ' * 3))
        document.add_picture(my_dir + 'img_costs_gross_profit_rate.png', width=pic_width)
        document.add_heading('2.三项费用率', level=2)
        document.add_paragraph(my_dict['three_fee_rate'].replace('\n', ' ' * 3))
        document.add_picture(my_dir + 'img_costs_three_fee_rate.png', width=pic_width)
        document.add_heading('3.销售费用率', level=2)
        document.add_paragraph(my_dict['sale_fee_rate'].replace('\n', ' ' * 3))
        document.add_picture(my_dir + 'img_costs_sale_fee_rate.png', width=pic_width)
        document.add_heading('4.管理费用率', level=2)
        document.add_paragraph(my_dict['manage_fee_rate'].replace('\n', ' ' * 3))
        document.add_picture(my_dir + 'img_costs_manage_fee_rate.png', width=pic_width)
        document.add_heading('5.财务费用率', level=2)
        document.add_paragraph(my_dict['finance_fee_rate'].replace('\n', ' ' * 3))
        document.add_picture(my_dir + 'img_costs_finance_fee_rate.png', width=pic_width)
    # 三.股东占比情况
    document.add_heading('三.股东占比情况', level=1)
    document.add_heading('1.股东人数', level=2)
    document.add_picture(my_dir + 'img_holder_count.png', width=pic_width)
    document.add_heading('2.十大流通股东', level=2)
    document.add_picture(my_dir + 'img_holder_ten_circulation.png', width=pic_width)
    document.add_heading('3.十大股东', level=2)
    document.add_picture(my_dir + 'img_holder_ten.png', width=pic_width)
    document.add_heading('4.控制层级关系', level=2)
    document.add_picture(my_dir + 'img_controller.png', width=pic_width)
    document.add_heading('5.机构持股汇总', level=2)
    document.add_picture(my_dir + 'img_position.png', width=pic_width)
    document.add_heading('6.高管持股变动', level=2)
    if os.path.isfile(my_dir + 'img_event_manager.png'):
        document.add_picture(my_dir + 'img_event_manager.png', width=pic_width)
    else:
        document.add_paragraph('无高管持股变动')
    if os.path.isfile(my_dir + 'img_event_holder.png'):
        document.add_heading('7.股东持股变动', level=2)
        document.add_picture(my_dir + 'img_event_holder.png', width=pic_width)
    document.add_heading('8.分红情况', level=2)
    document.add_picture(my_dir + 'img_bonus.png', width=pic_width)
    # 四.生意特征
    document.add_heading('四.生意特征', level=1)
    document.add_heading('1.主要客户及供应商', level=2)
    if os.path.isfile(my_dir + 'img_operate_partner.png'):
        document.add_picture(my_dir + 'img_operate_partner.png', width=pic_width)
    else:
        document.add_paragraph('无主要客户及供应商')
    document.add_heading('2.加权ROE', level=2)
    document.add_paragraph(my_dict['roe_weight'].replace('\n', ' ' * 3))
    document.add_picture(my_dir + 'img_profit_roe_weight.png', width=pic_width)
    document.add_heading('3.归属于母公司股东的ROE', level=2)
    document.add_paragraph(my_dict['roe'].replace('\n', ' ' * 3))
    document.add_picture(my_dir + 'img_profit_roe.png', width=pic_width)
    document.add_heading('4.杠杆倍数', level=2)
    document.add_paragraph(my_dict['leverage'].replace('\n', ' ' * 3))
    document.add_picture(my_dir + 'img_profit_leverage.png', width=pic_width)
    document.add_heading('5.资产周转率', level=2)
    document.add_paragraph(my_dict['turnover'].replace('\n', ' ' * 3))
    document.add_picture(my_dir + 'img_profit_turnover.png', width=pic_width)
    document.add_heading('6.净利润率', level=2)
    document.add_paragraph(my_dict['profit_rate'].replace('\n', ' ' * 3))
    document.add_picture(my_dir + 'img_profit_profit_rate.png', width=pic_width)
    # 五.成长性评估
    document.add_heading('五.成长性评估', level=1)
    document.add_heading('1.业绩预测', level=2)
    document.add_picture(my_dir + 'img_worth_forecast.png', width=pic_width)
    if os.path.isfile(my_dir + 'img_worth_forecast_table.png'):
        document.add_heading('2.业绩预测详表', level=2)
        document.add_picture(my_dir + 'img_worth_forecast_table.png', width=pic_width)
        document.add_picture(my_dir + 'img_worth_forecast_table2.png', width=pic_width)
    # 六.估值分析
    document.add_heading('六.估值分析', level=1)
    document.add_heading('1.PE-TTM (扣非)', level=2)
    document.add_paragraph(my_dict['pe'].replace('\n', ' ' * 3))
    document.add_picture(my_dir + 'img_valuation_pe.png', width=pic_width)
    document.add_heading('2.PB (不含商誉)', level=2)
    document.add_paragraph(my_dict['pb'].replace('\n', ' ' * 3))
    document.add_picture(my_dir + 'img_valuation_pb.png', width=pic_width)
    document.add_heading('3.PS-TTM', level=2)
    document.add_paragraph(my_dict['ps'].replace('\n', ' ' * 3))
    document.add_picture(my_dir + 'img_valuation_ps.png', width=pic_width)
    # 七.收入，利润，现金流分析
    document.add_heading('七.收入，利润，现金流分析', level=1)
    document.add_heading('1.营业总收入', level=2)
    document.add_paragraph(my_dict['total_income'].replace('\n', ' ' * 3))
    document.add_picture(my_dir + 'img_growth_total_income.png', width=pic_width)
    document.add_heading('2.营业利润', level=2)
    document.add_paragraph(my_dict['profit'].replace('\n', ' ' * 3))
    document.add_picture(my_dir + 'img_growth_profit.png', width=pic_width)
    document.add_heading('3.归属于母公司股东的扣非净利润', level=2)
    document.add_paragraph(my_dict['profit_exclude'].replace('\n', ' ' * 3))
    document.add_picture(my_dir + 'img_growth_profit_exclude.png', width=pic_width)
    document.add_heading('4.经营活动产生的现金流量净额', level=2)
    document.add_paragraph(my_dict['cash_flow'].replace('\n', ' ' * 3))
    document.add_picture(my_dir + 'img_growth_cash_flow.png', width=pic_width)
    # 八.资产负债分析
    document.add_heading('八.资产负债分析', level=1)
    if os.path.isfile(my_dir + 'img_asset.png') and os.path.isfile(my_dir + 'img_debt.png'):
        document.add_heading('1.资产分析', level=2)
        document.add_picture(my_dir + 'img_asset.png', width=pic_width)
        document.add_heading('2.负债分析', level=2)
        document.add_picture(my_dir + 'img_debt.png', width=pic_width)
    else:
        document.add_paragraph('特殊行业无资产负债分析')
    # 九.券商分析
    document.add_heading('九.券商分析', level=1)
    document.add_paragraph(util.read_file(code, 'file_worth'))
    # 十.题材要点
    document.add_heading('十.题材要点', level=1)
    document.add_paragraph(util.read_file(code, 'file_concept'))
    # 十一.董事会经营评述
    document.add_heading('十一.董事会经营评述', level=1)
    document.add_paragraph(util.read_file(code, 'file_operate'))
    # 保存文档
    document.save('private/' + code + '/' + my_dict['code_name'] + r'.docx')
    util.log('成功生成文档')


if __name__ == '__main__':
    generate_doc(config.CODE)
